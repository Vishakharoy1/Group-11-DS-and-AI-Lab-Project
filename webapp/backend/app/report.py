"""Builds forensic report documents for one analyzed image - prediction,
Grad-CAM evidence, methodology, and the reliability caveats established
in the Milestone 5 evaluation (real numbers, not generic disclaimers).

build_report_html() is the legacy standalone HTML report (still callable
via POST /report). build_docx() backs the new frontend's "Download Report
-> Word Document (.docx)" action - the frontend Report page itself is the
canonical on-screen rendering; this just needs to reproduce the same data
as a downloadable .docx."""

import base64
import io
import uuid
from datetime import datetime, timezone

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FACE_ALIGNMENT_LABELS = {
    "retinaface": "RetinaFace (automatic face detection & alignment)",
    "center_crop_fallback": "Center-crop fallback (RetinaFace unavailable on this server)",
}

INFERENCE_STEPS_APPLIED = ["Face Detection", "Image Resize", "Normalization", "Tensor Conversion"]
TRAINING_AUGMENTATIONS = [
    "Random Resized Crop",
    "Random Horizontal Flip",
    "Color Jitter",
    "Channel Shift",
    "Gaussian Blur",
    "JPEG Compression",
    "Gaussian Noise",
]

RELIABILITY_NOTES = [
    "Strong on in-distribution data: 99.63% accuracy on the model's own held-out test set.",
    "Weak on modern smartphone photos: accuracy on genuine recent smartphone photographs dropped to "
    "8.6% in cross-domain testing, and 62.0% in a separate local check. If this input is a modern "
    "smartphone photo (HDR, heavy sharpening, high saturation), treat an AI-Generated verdict with caution.",
    "Dominant error type is false positives: genuine photos being misclassified as AI-generated, not "
    "fakes evading detection.",
    "Non-face images produce meaningless results: if the input does not contain a clearly detectable "
    "face, this verdict should be disregarded entirely.",
]

DISCLAIMER_TEXT = (
    "This report represents the output of an AI-based image authenticity detection model and should be "
    "interpreted as an analytical assessment rather than definitive proof of image manipulation or "
    "authenticity. It does not constitute a certified, legally admissible forensic conclusion. No "
    "chain-of-custody or examiner certification is implied."
)


def build_report_html(
    *,
    input_image_b64: str,
    filename: str,
    model_name: str,
    face_alignment_used: str,
    label: str,
    real_pct: float,
    fake_pct: float,
    heatmap_b64: str,
    overlay_b64: str,
) -> str:
    report_id = str(uuid.uuid4())[:8].upper()
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    alignment_label = FACE_ALIGNMENT_LABELS.get(face_alignment_used, face_alignment_used)
    verdict_class = "real" if label == "Real" else "fake"
    confidence = real_pct if label == "Real" else fake_pct

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8" />
<title>Face Authenticity Report — {report_id}</title>
<style>
  @media print {{
    body {{ margin: 0; }}
    .no-print {{ display: none; }}
  }}
  body {{
    font-family: Georgia, 'Times New Roman', serif;
    max-width: 800px;
    margin: 40px auto;
    padding: 0 24px;
    color: #1a1a1a;
    line-height: 1.55;
  }}
  header {{
    border-bottom: 3px solid #1a1a1a;
    padding-bottom: 16px;
    margin-bottom: 24px;
  }}
  h1 {{ font-size: 1.5rem; margin: 0 0 4px 0; }}
  .meta {{ font-family: 'Courier New', monospace; font-size: 0.85rem; color: #444; }}
  .meta span {{ display: inline-block; margin-right: 24px; }}
  section {{ margin-bottom: 28px; page-break-inside: avoid; }}
  h2 {{
    font-size: 1.05rem;
    text-transform: uppercase;
    letter-spacing: 0.04em;
    border-bottom: 1px solid #ccc;
    padding-bottom: 6px;
    margin-bottom: 14px;
  }}
  .verdict-box {{
    display: flex;
    align-items: center;
    gap: 20px;
    padding: 18px;
    border: 2px solid;
    border-radius: 6px;
  }}
  .verdict-box.real {{ border-color: #2e7d32; background: #f1f8f2; }}
  .verdict-box.fake {{ border-color: #c62828; background: #fdf1f1; }}
  .verdict-label {{ font-size: 1.6rem; font-weight: bold; }}
  .verdict-box.real .verdict-label {{ color: #2e7d32; }}
  .verdict-box.fake .verdict-label {{ color: #c62828; }}
  .pct-row {{ font-family: 'Courier New', monospace; font-size: 0.95rem; margin-top: 4px; }}
  .images-grid {{
    display: flex;
    gap: 16px;
    flex-wrap: wrap;
  }}
  .images-grid figure {{ margin: 0; text-align: center; }}
  .images-grid img {{
    width: 200px;
    height: 200px;
    object-fit: cover;
    border: 1px solid #ccc;
    border-radius: 4px;
  }}
  figcaption {{ font-size: 0.8rem; color: #555; margin-top: 6px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 0.9rem; }}
  table td {{ padding: 6px 0; border-bottom: 1px solid #eee; }}
  table td:first-child {{ color: #555; width: 220px; }}
  .caveats {{
    background: #fff8e1;
    border-left: 4px solid #f9a825;
    padding: 14px 18px;
    font-size: 0.9rem;
  }}
  .caveats ul {{ margin: 8px 0 0 0; padding-left: 20px; }}
  .caveats li {{ margin-bottom: 6px; }}
  .disclaimer {{
    font-size: 0.8rem;
    color: #666;
    border-top: 1px solid #ccc;
    padding-top: 14px;
    margin-top: 32px;
  }}
  .print-btn {{
    background: #1a1a1a;
    color: white;
    border: none;
    padding: 10px 20px;
    border-radius: 4px;
    cursor: pointer;
    font-size: 0.9rem;
    margin-bottom: 20px;
  }}
</style>
</head>
<body>

<button class="print-btn no-print" onclick="window.print()">Print / Save as PDF</button>

<header>
  <h1>Face Authenticity Analysis Report</h1>
  <div class="meta">
    <span>Report ID: {report_id}</span>
    <span>Generated: {generated_at}</span>
  </div>
</header>

<section>
  <h2>1. Verdict</h2>
  <div class="verdict-box {verdict_class}">
    <div>
      <div class="verdict-label">{label}</div>
      <div class="pct-row">Confidence: {confidence:.2f}%</div>
      <div class="pct-row">Real: {real_pct:.2f}% &nbsp;|&nbsp; Fake: {fake_pct:.2f}%</div>
    </div>
  </div>
</section>

<section>
  <h2>2. Case Details</h2>
  <table>
    <tr><td>Input filename</td><td>{filename}</td></tr>
    <tr><td>Model checkpoint used</td><td><code>{model_name}</code></td></tr>
    <tr><td>Face alignment method</td><td>{alignment_label}</td></tr>
  </table>
</section>

<section>
  <h2>3. Input &amp; Explainability Evidence</h2>
  <div class="images-grid">
    <figure>
      <img src="data:image/png;base64,{input_image_b64}" alt="Input image" />
      <figcaption>Input (as analyzed)</figcaption>
    </figure>
    <figure>
      <img src="data:image/png;base64,{heatmap_b64}" alt="Grad-CAM heatmap" />
      <figcaption>Grad-CAM heatmap</figcaption>
    </figure>
    <figure>
      <img src="data:image/png;base64,{overlay_b64}" alt="Grad-CAM overlay" />
      <figcaption>Grad-CAM overlay ({label})</figcaption>
    </figure>
  </div>
  <p style="font-size:0.85rem; color:#555; margin-top:12px;">
    The heatmap highlights which regions of the image most influenced the
    model's decision — brighter regions contributed more. A heatmap
    concentrated on facial structure (eyes, nose, mouth) is consistent
    with the model attending to genuine facial evidence; a heatmap
    concentrated on background or edges may indicate a less reliable
    prediction.
  </p>
</section>

<section>
  <h2>4. Reliability &amp; Known Limitations</h2>
  <div class="caveats">
    This model was evaluated in detail as part of a separate technical
    milestone report. Key findings relevant to interpreting this result:
    <ul>
      <li><strong>Strong on in-distribution data:</strong> 99.63% accuracy on the model's own held-out test set.</li>
      <li><strong>Weak on modern smartphone photos:</strong> accuracy on genuine recent smartphone photographs dropped to 8.6% in cross-domain testing, and 62.0% in a separate local check — both substantially below the in-distribution figure. If this input is a modern smartphone photo (HDR, heavy sharpening, high saturation), treat a "Fake" verdict with caution.</li>
      <li><strong>Dominant error type is false positives:</strong> genuine photos being misclassified as fake, not fakes evading detection.</li>
      <li><strong>Non-face images produce meaningless results:</strong> if the input does not contain a clearly detectable face, this verdict should be disregarded entirely.</li>
    </ul>
  </div>
</section>

<div class="disclaimer">
  This is an automated, educational analysis tool. It does not constitute
  a certified, legally admissible forensic conclusion. No chain-of-custody
  or examiner certification is implied. Results should be interpreted
  alongside the limitations stated above, not as a standalone determination
  of authenticity.
</div>

</body>
</html>
"""


def _b64_to_stream(b64_str: str) -> io.BytesIO:
    return io.BytesIO(base64.b64decode(b64_str))


def build_docx(
    *,
    analysis_id: str,
    generated_at_str: str,
    filename: str,
    file_type: str,
    resolution: str,
    file_size: str,
    color_mode: str,
    model_label: str,
    model_version: str,
    label: str,
    real_pct: float,
    fake_pct: float,
    input_image_b64: str,
    overlay_b64: str,
) -> io.BytesIO:
    """Builds a .docx matching the Report page's on-screen sections. Returns
    an in-memory buffer ready to stream as a download."""
    doc = Document()

    verdict_text = "AI GENERATED" if label != "Real" else "REAL / AUTHENTIC"
    verdict_color = RGBColor(0xC6, 0x28, 0x28) if label != "Real" else RGBColor(0x2E, 0x7D, 0x32)
    confidence = fake_pct if label != "Real" else real_pct

    title = doc.add_heading("AI Image Authenticity", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = doc.add_heading("Forensic Analysis Report", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Analysis ID: {analysis_id}    |    Generated: {generated_at_str}").italic = True

    doc.add_heading("1. Case Information", level=2)
    table1 = doc.add_table(rows=0, cols=2)
    for k, v in [("Analysis ID", analysis_id), ("Date & Time", generated_at_str), ("Uploaded Filename", filename)]:
        row = table1.add_row().cells
        row[0].text, row[1].text = k, v

    doc.add_heading("2. Image Information", level=2)
    table2 = doc.add_table(rows=0, cols=2)
    for k, v in [
        ("File Type", file_type),
        ("Resolution", resolution),
        ("File Size", file_size),
        ("Color Mode", color_mode),
    ]:
        row = table2.add_row().cells
        row[0].text, row[1].text = k, v

    doc.add_heading("3. Model Information", level=2)
    table3 = doc.add_table(rows=0, cols=2)
    for k, v in [
        ("Model", model_label),
        ("Model Version", model_version),
        ("Prediction", verdict_text),
        ("Confidence", f"{confidence:.2f}%"),
        ("AI Probability", f"{fake_pct:.2f}%"),
        ("Real Probability", f"{real_pct:.2f}%"),
    ]:
        row = table3.add_row().cells
        row[0].text, row[1].text = k, v

    doc.add_heading("4. Inference Preprocessing Pipeline", level=2)
    p = doc.add_paragraph("Applied at inference time:")
    for step in INFERENCE_STEPS_APPLIED:
        doc.add_paragraph(step, style="List Bullet")
    p2 = doc.add_paragraph("Training augmentation (used during model training, NOT applied during inference):")
    for aug in TRAINING_AUGMENTATIONS:
        doc.add_paragraph(aug, style="List Bullet")

    doc.add_heading("5. Explainability Analysis (Grad-CAM)", level=2)
    img_table = doc.add_table(rows=1, cols=2)
    input_cell, overlay_cell = img_table.rows[0].cells
    input_cell.paragraphs[0].add_run().add_picture(_b64_to_stream(input_image_b64), width=Inches(2.6))
    overlay_cell.paragraphs[0].add_run().add_picture(_b64_to_stream(overlay_b64), width=Inches(2.6))
    doc.add_paragraph(
        "The Grad-CAM overlay highlights regions that received stronger activation during the model's "
        "prediction. These regions indicate where the model focused its attention when producing the "
        "classification - this shows model attention, not definitive proof of manipulation or authenticity."
    )

    doc.add_heading("6. Final Assessment", level=2)
    verdict_para = doc.add_paragraph()
    verdict_run = verdict_para.add_run(f"{verdict_text} — {confidence:.2f}%")
    verdict_run.bold = True
    verdict_run.font.size = Pt(16)
    verdict_run.font.color.rgb = verdict_color
    doc.add_paragraph(
        f"The model classified the submitted image as {verdict_text.lower()} with a confidence score of "
        f"{confidence:.2f}%."
    )

    doc.add_heading("Reliability & Known Limitations", level=2)
    for note in RELIABILITY_NOTES:
        doc.add_paragraph(note, style="List Bullet")

    doc.add_heading("7. Disclaimer", level=2)
    disclaimer_para = doc.add_paragraph(DISCLAIMER_TEXT)
    disclaimer_para.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
